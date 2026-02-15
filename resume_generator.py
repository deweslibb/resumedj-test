#!/usr/bin/env python3
"""
Resume Generator - Fully Dynamic ATS-Friendly Version with Custom Fonts & Sections
Reads section order from Excel, handles all sections dynamically with user customization
"""
import sys
from pathlib import Path
from datetime import datetime
from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT

def safe_str(value):
    """Safely convert cell value to string"""
    if value is None or value == "":
        return ""
    return str(value).strip()

def default_include(value):
    """Default to 'No' if blank"""
    val = safe_str(value).upper()
    if val in ["YES", "Y"]:
        return "Yes"
    return "No"

def default_order(value):
    """Default to 9999 (last) if blank"""
    if value is None or safe_str(value) == "":
        return 9999
    try:
        return int(value)
    except:
        return 9999

def strip_bullet_prefix(text):
    """Remove bullet characters and leading spaces from text"""
    if not text:
        return text
    
    # Common bullet characters to remove
    bullets = ['•', '●', '○', '◦', '▪', '▫', '■', '□', '-', '*', '→', '⇒', '►', '▸']
    result = text.strip()
    
    # Remove bullet if it's the first character
    if result and result[0] in bullets:
        result = result[1:].strip()
    
    return result

def format_date(value):
    """Format date values"""
    val = safe_str(value).upper()
    if val == "PRESENT":
        return "Present"
    if value and hasattr(value, 'strftime'):
        return value.strftime("%b %Y")
    return safe_str(value)

def read_defaults(wb):
    """Read font defaults from Defaults sheet"""
    defaults = {
        'Name & Credentials': {'font': 'Calibri', 'size': 20, 'bold': True, 'italic': False},
        'Title': {'font': 'Calibri', 'size': 14, 'bold': True, 'italic': False},
        'Contact Information': {'font': 'Calibri', 'size': 10, 'bold': False, 'italic': False},
        'Section Headers': {'font': 'Calibri', 'size': 11, 'bold': True, 'italic': False},
        'Company/School Names': {'font': 'Calibri', 'size': 10, 'bold': True, 'italic': False},
        'Employed/Enrolled Dates': {'font': 'Calibri', 'size': 10, 'bold': False, 'italic': False},
        'Job Title/Degree': {'font': 'Calibri', 'size': 10, 'bold': False, 'italic': True},
        'Location': {'font': 'Calibri', 'size': 10, 'bold': False, 'italic': True},
        'Body (Bullets)': {'font': 'Calibri', 'size': 10, 'bold': False, 'italic': False}
    }
    
    try:
        ws = wb['Defaults']
        for row in range(3, 12):  # Rows 3-11 contain the settings
            field = safe_str(ws.cell(row, 2).value)
            font = safe_str(ws.cell(row, 3).value)
            size = ws.cell(row, 4).value
            bold = safe_str(ws.cell(row, 5).value).upper()
            italic = safe_str(ws.cell(row, 6).value).upper()
            
            if field in defaults:
                # Only update if font is provided
                if font:
                    defaults[field]['font'] = font
                if size:
                    try:
                        defaults[field]['size'] = int(size)
                    except:
                        pass  # Keep default
                defaults[field]['bold'] = (bold == 'YES')
                defaults[field]['italic'] = (italic == 'YES')
    except:
        pass  # Return defaults if sheet doesn't exist or error
    
    return defaults

def read_custom_section_name(ws):
    """Read custom section name from B2 of a sheet"""
    try:
        custom_name = safe_str(ws.cell(2, 2).value)
        # Remove instruction text if present
        if '<----' in custom_name:
            custom_name = custom_name.split('<----')[0].strip()
        return custom_name if custom_name else None
    except:
        return None

def read_section_order(wb):
    """Read which sections to include and in what order"""
    try:
        ws = wb['Order']
    except:
        # Default order if Order sheet doesn't exist
        return [
            'Contact_Info', 'Summary', 'Highlights', 
            'Work_Roles', 'Education', 
            'Licenses_Certifications', 'Skills', 'Achievements'
        ]
    
    sections = []
    for row in range(2, 50):  # Start at row 2, skip header
        section_name = safe_str(ws.cell(row, 3).value)  # Column C
        include = default_include(ws.cell(row, 4).value)  # Column D
        
        if not section_name:
            break
        
        if include == "Yes":
            sections.append(section_name)
    
    return sections

def read_contact_info(ws):
    """Read contact information with new structure"""
    contact = {}
    for row in range(4, 15):
        field = safe_str(ws.cell(row, 2).value)
        value = safe_str(ws.cell(row, 3).value)
        include = default_include(ws.cell(row, 4).value)
        
        if include == "Yes" and value:
            contact[field] = value
    
    return contact

def read_summary(ws):
    """Read summary text"""
    include = default_include(ws.cell(3, 4).value)  # Row 3, Column D
    summary_text = safe_str(ws.cell(5, 2).value)  # Row 5, Column B
    
    if include == "Yes" and summary_text:
        return summary_text
    return None

def read_work_roles(ws):
    """Read work roles with Include? filter"""
    roles = []
    for row in range(4, 101):
        company = safe_str(ws.cell(row, 2).value)
        if not company:
            break
        
        include = default_include(ws.cell(row, 7).value)
        
        if include == "Yes":
            roles.append({
                'company': company,
                'title': safe_str(ws.cell(row, 3).value),
                'start_date': format_date(ws.cell(row, 4).value),
                'end_date': format_date(ws.cell(row, 5).value),
                'location': safe_str(ws.cell(row, 6).value)
            })
    
    return roles

def read_work_experience(ws, company):
    """Read and sort bullets for a specific company"""
    bullets = []
    
    for row in range(4, 201):
        comp = safe_str(ws.cell(row, 2).value)
        if comp != company:
            continue
        
        bullet_text = safe_str(ws.cell(row, 3).value)
        include = default_include(ws.cell(row, 4).value)
        order = default_order(ws.cell(row, 5).value)
        
        if include == "Yes" and bullet_text:
            # Strip bullet characters and leading spaces
            cleaned_text = strip_bullet_prefix(bullet_text)
            bullets.append((order, cleaned_text))
    
    bullets.sort(key=lambda x: x[0])
    return [b[1] for b in bullets]

def read_bullet_section(ws):
    """Generic function to read bullet sections"""
    bullets = []
    
    for row in range(4, 201):
        bullet_text = safe_str(ws.cell(row, 2).value)  # Column B
        include = default_include(ws.cell(row, 3).value)  # Column C
        order = default_order(ws.cell(row, 4).value)  # Column D
        
        if not bullet_text:
            break
        
        if include == "Yes" and bullet_text:
            bullets.append((order, bullet_text))
    
    bullets.sort(key=lambda x: x[0])
    return [b[1] for b in bullets]

def read_education(ws):
    """Read education information"""
    education = []
    for row in range(4, 51):
        school = safe_str(ws.cell(row, 2).value)
        if not school:
            break
        
        education.append({
            'school': school,
            'degree': safe_str(ws.cell(row, 3).value),
            'major': safe_str(ws.cell(row, 4).value),
            'end_date': format_date(ws.cell(row, 6).value),
            'location': safe_str(ws.cell(row, 7).value),
            'graduated': safe_str(ws.cell(row, 8).value)
        })
    
    return education

def read_skills(ws, col_idx):
    """Read and sort skills for a category"""
    skills = []
    
    for row in range(4, 51):
        skill = safe_str(ws.cell(row, col_idx).value)
        include = default_include(ws.cell(row, col_idx + 1).value)
        order = default_order(ws.cell(row, col_idx + 2).value)
        
        if include == "Yes" and skill:
            skills.append((order, skill))
    
    skills.sort(key=lambda x: x[0])
    return [s[1] for s in skills]

def safe_font_apply(run, font_name):
    """Safely apply font with fallback to Times New Roman"""
    try:
        run.font.name = font_name
    except:
        try:
            run.font.name = 'Times New Roman'
        except:
            pass  # Keep default if both fail

def build_contact_header(doc, contact, defaults):
    """Build contact header: Name+Credentials, optional Title, then other info"""
    # Line 1: Name + Credentials
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    name_line = contact.get('Name', '')
    if 'Credentials' in contact:
        name_line += f", {contact['Credentials']}"
    
    run = p.add_run(name_line)
    name_style = defaults.get('Name & Credentials', {})
    safe_font_apply(run, name_style.get('font', 'Times New Roman'))
    run.font.size = Pt(name_style.get('size', 20))
    run.bold = name_style.get('bold', True)
    run.italic = name_style.get('italic', False)
    p.space_after = Pt(4)
    
    # Line 2: Title (optional)
    if 'Title' in contact:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(contact['Title'])
        title_style = defaults.get('Title', {})
        safe_font_apply(run, title_style.get('font', 'Times New Roman'))
        run.font.size = Pt(title_style.get('size', 14))
        run.bold = title_style.get('bold', True)
        run.italic = title_style.get('italic', False)
        p.space_after = Pt(4)
    
    # Line 3: Everything else on one line
    contact_parts = []
    for field in ['Email', 'Phone', 'LinkedIn', 'Website']:
        if field in contact:
            contact_parts.append(contact[field])
    
    # Add location
    location_parts = []
    if 'City' in contact:
        location_parts.append(contact['City'])
    if 'State' in contact:
        location_parts.append(contact['State'])
    if 'Country' in contact:
        location_parts.append(contact['Country'])
    if location_parts:
        contact_parts.append(', '.join(location_parts))
    
    if contact_parts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(' | '.join(contact_parts))
        contact_style = defaults.get('Contact Information', {})
        safe_font_apply(run, contact_style.get('font', 'Times New Roman'))
        run.font.size = Pt(contact_style.get('size', 10))
        run.bold = contact_style.get('bold', False)
        run.italic = contact_style.get('italic', False)
        p.space_after = Pt(6)

def build_summary(doc, summary_text, defaults, custom_name=None):
    """Build summary section"""
    if not summary_text:
        return
    
    # Section header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    section_name = custom_name if custom_name else 'Professional Summary'
    run = p.add_run(section_name.upper())
    header_style = defaults.get('Section Headers', {})
    safe_font_apply(run, header_style.get('font', 'Times New Roman'))
    run.font.size = Pt(header_style.get('size', 11))
    run.bold = header_style.get('bold', True)
    run.italic = header_style.get('italic', False)
    p.space_after = Pt(4)
    
    # Summary text
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(summary_text)
    body_style = defaults.get('Body (Bullets)', {})
    safe_font_apply(run, body_style.get('font', 'Times New Roman'))
    run.font.size = Pt(body_style.get('size', 10))
    run.bold = body_style.get('bold', False)
    run.italic = body_style.get('italic', False)
    p.space_after = Pt(6)

def build_bullet_section(doc, title, bullets, defaults):
    """Build generic bullet point section"""
    if not bullets:
        return
    
    # Section header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(title.upper())
    header_style = defaults.get('Section Headers', {})
    safe_font_apply(run, header_style.get('font', 'Times New Roman'))
    run.font.size = Pt(header_style.get('size', 11))
    run.bold = header_style.get('bold', True)
    run.italic = header_style.get('italic', False)
    p.space_after = Pt(4)
    
    # Bullets
    for bullet in bullets:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(bullet)
        body_style = defaults.get('Body (Bullets)', {})
        safe_font_apply(run, body_style.get('font', 'Times New Roman'))
        run.font.size = Pt(body_style.get('size', 10))
        run.bold = body_style.get('bold', False)
        run.italic = body_style.get('italic', False)
    
    # Add spacing after section
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.space_after = Pt(6)

def build_work_experience(doc, roles, work_exp_ws, defaults, custom_name=None):
    """Build work experience section"""
    if not roles:
        return
    
    # Section header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    section_name = custom_name if custom_name else 'PROFESSIONAL EXPERIENCE'
    run = p.add_run(section_name.upper())
    header_style = defaults.get('Section Headers', {})
    safe_font_apply(run, header_style.get('font', 'Times New Roman'))
    run.font.size = Pt(header_style.get('size', 11))
    run.bold = header_style.get('bold', True)
    run.italic = header_style.get('italic', False)
    p.space_after = Pt(4)
    
    # Each role
    for role in roles:
        # Company and dates on one line with tab stop for right-alignment
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(0)
        
        # Set tab stop at 7.5 inches for right-alignment
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)
        
        # Company name (left)
        run = p.add_run(role['company'])
        company_style = defaults.get('Company/School Names', {})
        safe_font_apply(run, company_style.get('font', 'Times New Roman'))
        run.font.size = Pt(company_style.get('size', 10))
        run.bold = company_style.get('bold', True)
        run.italic = company_style.get('italic', False)
        
        # Tab + Dates (right-aligned)
        date_range = f"{role['start_date']} - {role['end_date']}"
        run = p.add_run('\t' + date_range)
        date_style = defaults.get('Employed/Enrolled Dates', {})
        safe_font_apply(run, date_style.get('font', 'Times New Roman'))
        run.font.size = Pt(date_style.get('size', 10))
        run.bold = date_style.get('bold', False)
        run.italic = date_style.get('italic', False)
        
        # Title and location on next line
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        
        # Set tab stop at 7.5 inches for right-alignment
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)
        
        # Title (left)
        run = p.add_run(role['title'])
        title_style = defaults.get('Job Title/Degree', {})
        safe_font_apply(run, title_style.get('font', 'Times New Roman'))
        run.font.size = Pt(title_style.get('size', 10))
        run.bold = title_style.get('bold', False)
        run.italic = title_style.get('italic', True)
        
        # Tab + Location (right-aligned)
        if role.get('location'):
            run = p.add_run('\t' + role['location'])
            loc_style = defaults.get('Location', {})
            safe_font_apply(run, loc_style.get('font', 'Times New Roman'))
            run.font.size = Pt(loc_style.get('size', 10))
            run.bold = loc_style.get('bold', False)
            run.italic = loc_style.get('italic', True)
        
        # Bullets for this company
        bullets = read_work_experience(work_exp_ws, role['company'])
        for bullet in bullets:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(bullet)
            body_style = defaults.get('Body (Bullets)', {})
            safe_font_apply(run, body_style.get('font', 'Times New Roman'))
            run.font.size = Pt(body_style.get('size', 10))
            run.bold = body_style.get('bold', False)
            run.italic = body_style.get('italic', False)
        
        # Extra space between jobs
        if role != roles[-1]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
    
    # Add spacing after section
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.space_after = Pt(6)

def build_education(doc, education, defaults, custom_name=None):
    """Build education section"""
    if not education:
        return
    
    # Section header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    section_name = custom_name if custom_name else 'EDUCATION'
    run = p.add_run(section_name.upper())
    header_style = defaults.get('Section Headers', {})
    safe_font_apply(run, header_style.get('font', 'Times New Roman'))
    run.font.size = Pt(header_style.get('size', 11))
    run.bold = header_style.get('bold', True)
    run.italic = header_style.get('italic', False)
    p.space_after = Pt(4)
    
    # Each degree
    for edu in education:
        # School and date on one line with tab stop
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(0)
        
        # Set tab stop at 7.5 inches for right-alignment
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)
        
        # School name (left)
        run = p.add_run(edu['school'])
        school_style = defaults.get('Company/School Names', {})
        safe_font_apply(run, school_style.get('font', 'Times New Roman'))
        run.font.size = Pt(school_style.get('size', 10))
        run.bold = school_style.get('bold', True)
        run.italic = school_style.get('italic', False)
        
        # Graduation status and date (right-aligned)
        if edu.get('graduated', '').upper() == 'YES':
            date_text = f"Graduated {edu.get('end_date', '')}"
        elif edu.get('graduated', '').upper() == 'NO':
            if edu.get('end_date'):
                date_text = f"Anticipated {edu.get('end_date')}"
            else:
                date_text = "In Progress"
        else:
            date_text = edu.get('end_date', '')
        
        if date_text:
            run = p.add_run('\t' + date_text)
            date_style = defaults.get('Employed/Enrolled Dates', {})
            safe_font_apply(run, date_style.get('font', 'Times New Roman'))
            run.font.size = Pt(date_style.get('size', 10))
            run.bold = date_style.get('bold', False)
            run.italic = date_style.get('italic', False)
        
        # Degree and location on next line
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        
        # Set tab stop at 7.5 inches for right-alignment
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)
        
        # Degree and major
        degree_text = edu.get('degree', '')
        if edu.get('major'):
            degree_text += f", {edu['major']}"
        
        run = p.add_run(degree_text)
        degree_style = defaults.get('Job Title/Degree', {})
        safe_font_apply(run, degree_style.get('font', 'Times New Roman'))
        run.font.size = Pt(degree_style.get('size', 10))
        run.bold = degree_style.get('bold', False)
        run.italic = degree_style.get('italic', True)
        
        # Location (right-aligned)
        if edu.get('location'):
            run = p.add_run('\t' + edu['location'])
            loc_style = defaults.get('Location', {})
            safe_font_apply(run, loc_style.get('font', 'Times New Roman'))
            run.font.size = Pt(loc_style.get('size', 10))
            run.bold = loc_style.get('bold', False)
            run.italic = loc_style.get('italic', True)
    
    # Add spacing after section
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.space_after = Pt(6)

def build_skills(doc, skills_ws, defaults, custom_name=None):
    """Build skills section with up to 4 categories"""
    # Read category names and skills
    categories = []
    
    for col_idx in [2, 6, 10, 14]:  # Columns B, F, J, N
        category_name = safe_str(skills_ws.cell(3, col_idx).value)
        if not category_name:
            continue
        
        skills = read_skills(skills_ws, col_idx)
        if skills:
            categories.append({
                'name': category_name,
                'skills': skills
            })
    
    if not categories:
        return
    
    # Section header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    section_name = custom_name if custom_name else 'SKILLS'
    run = p.add_run(section_name.upper())
    header_style = defaults.get('Section Headers', {})
    safe_font_apply(run, header_style.get('font', 'Times New Roman'))
    run.font.size = Pt(header_style.get('size', 11))
    run.bold = header_style.get('bold', True)
    run.italic = header_style.get('italic', False)
    p.space_after = Pt(4)
    
    # Each category
    for cat in categories:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        
        # Category name in bold
        run = p.add_run(f"{cat['name']}: ")
        body_style = defaults.get('Body (Bullets)', {})
        safe_font_apply(run, body_style.get('font', 'Times New Roman'))
        run.font.size = Pt(body_style.get('size', 10))
        run.bold = True
        run.italic = body_style.get('italic', False)
        
        # Skills comma-separated
        run = p.add_run(', '.join(cat['skills']))
        safe_font_apply(run, body_style.get('font', 'Times New Roman'))
        run.font.size = Pt(body_style.get('size', 10))
        run.bold = body_style.get('bold', False)
        run.italic = body_style.get('italic', False)
    
    # Add spacing after section
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.space_after = Pt(6)

def create_resume(excel_path):
    """Main function to create resume"""
    print(f"Loading Excel file: {excel_path}")
    wb = load_workbook(excel_path, data_only=True)
    
    # Read defaults
    defaults = read_defaults(wb)
    
    print("Creating Word document...")
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    try:
        style.font.name = 'Times New Roman'
    except:
        pass
    style.font.size = Pt(10)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0
    
    # Page setup
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    
    # Read section order
    section_order = read_section_order(wb)
    print(f"Section order: {section_order}")
    
    # Pre-read all data
    contact = read_contact_info(wb['Contact_Info'])
    
    summary = None
    summary_name = 'Professional Summary'
    if 'Summary' in wb.sheetnames:
        summary = read_summary(wb['Summary'])
        custom = read_custom_section_name(wb['Summary'])
        if custom:
            summary_name = custom
    
    highlights = []
    highlights_name = 'Highlights'
    if 'Highlights' in wb.sheetnames:
        highlights = read_bullet_section(wb['Highlights'])
        custom = read_custom_section_name(wb['Highlights'])
        if custom:
            highlights_name = custom
    
    achievements = []
    achievements_name = 'Achievements'
    if 'Achievements' in wb.sheetnames:
        achievements = read_bullet_section(wb['Achievements'])
        custom = read_custom_section_name(wb['Achievements'])
        if custom:
            achievements_name = custom
    
    licenses = []
    licenses_name = 'Licenses & Certifications'
    if 'Licenses_Certifications' in wb.sheetnames:
        licenses = read_bullet_section(wb['Licenses_Certifications'])
        custom = read_custom_section_name(wb['Licenses_Certifications'])
        if custom:
            licenses_name = custom
    
    memberships = []
    memberships_name = 'Memberships'
    if 'Memberships' in wb.sheetnames:
        memberships = read_bullet_section(wb['Memberships'])
        custom = read_custom_section_name(wb['Memberships'])
        if custom:
            memberships_name = custom
    
    awards = []
    awards_name = 'Awards'
    if 'Awards' in wb.sheetnames:
        awards = read_bullet_section(wb['Awards'])
        custom = read_custom_section_name(wb['Awards'])
        if custom:
            awards_name = custom
    
    projects = []
    projects_name = 'Projects'
    if 'Projects' in wb.sheetnames:
        projects = read_bullet_section(wb['Projects'])
        custom = read_custom_section_name(wb['Projects'])
        if custom:
            projects_name = custom
    
    custom_section = []
    custom_section_name = 'Additional Information'
    if 'Custom_Section' in wb.sheetnames:
        custom_section = read_bullet_section(wb['Custom_Section'])
        custom = read_custom_section_name(wb['Custom_Section'])
        if custom:
            custom_section_name = custom
    
    skills_name = 'Skills'
    if 'Skills' in wb.sheetnames:
        custom = read_custom_section_name(wb['Skills'])
        if custom:
            skills_name = custom
    
    work_roles_name = 'Professional Experience'
    if 'Work_Roles' in wb.sheetnames:
        custom = read_custom_section_name(wb['Work_Roles'])
        if custom:
            work_roles_name = custom
    
    work_roles = read_work_roles(wb['Work_Roles'])
    
    education_name = 'Education'
    if 'Education' in wb.sheetnames:
        custom = read_custom_section_name(wb['Education'])
        if custom:
            education_name = custom
    
    education = read_education(wb['Education'])
    
    # Build resume in the order specified by the Order sheet
    for section_name in section_order:
        print(f"Building section: {section_name}")
        
        if section_name == 'Contact_Info':
            build_contact_header(doc, contact, defaults)
        
        elif section_name == 'Summary':
            build_summary(doc, summary, defaults, summary_name)
        
        elif section_name == 'Highlights':
            build_bullet_section(doc, highlights_name, highlights, defaults)
        
        elif section_name == 'Work_Roles':
            build_work_experience(doc, work_roles, wb['Work_Experience'], defaults, work_roles_name)
        
        elif section_name == 'Education':
            build_education(doc, education, defaults, education_name)
        
        elif section_name == 'Licenses_Certifications':
            build_bullet_section(doc, licenses_name, licenses, defaults)
        
        elif section_name == 'Skills':
            build_skills(doc, wb['Skills'], defaults, skills_name)
        
        elif section_name == 'Achievements':
            build_bullet_section(doc, achievements_name, achievements, defaults)
        
        elif section_name == 'Memberships':
            build_bullet_section(doc, memberships_name, memberships, defaults)
        
        elif section_name == 'Awards':
            build_bullet_section(doc, awards_name, awards, defaults)
        
        elif section_name == 'Projects':
            build_bullet_section(doc, projects_name, projects, defaults)
        
        elif section_name == 'Custom_Section':
            build_bullet_section(doc, custom_section_name, custom_section, defaults)
    
    # Get user's Documents folder
    import os
    if os.name == 'nt':  # Windows
        documents_path = Path(os.path.expanduser("~")) / "Documents"
    else:  # Mac/Linux
        documents_path = Path(os.path.expanduser("~")) / "Documents"
    
    # Create Documents folder if it doesn't exist
    documents_path.mkdir(parents=True, exist_ok=True)
    
    # Create filename from first and last name
    full_name = contact.get('Name', 'Resume')
    name_parts = full_name.split()
    
    if len(name_parts) >= 2:
        first_name = name_parts[0]
        last_name = name_parts[-1]
        base_filename = f"{first_name}_{last_name}"
    else:
        base_filename = full_name.replace(' ', '_') if full_name else "Resume"
    
    # Add timestamp
    timestamp = datetime.now().strftime("%Y-%m-%d_%H%M%S")
    filename = f"{base_filename}_{timestamp}"
    
    # Full paths
    docx_path = documents_path / f"{filename}.docx"
    pdf_path = documents_path / f"{filename}.pdf"
    
    print(f"Saving to: {documents_path}")
    print(f"Filename: {filename}")
    print(f"Saving Word document: {docx_path}")
    doc.save(str(docx_path))
    
    print(f"Converting to PDF: {pdf_path}")
    
    try:
        import win32com.client
        import pythoncom
        
        pythoncom.CoInitialize()
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc_com = word.Documents.Open(str(docx_path.absolute()))
        doc_com.SaveAs(str(pdf_path.absolute()), FileFormat=17)
        doc_com.Close()
        word.Quit()
        pythoncom.CoUninitialize()
        
        print("✓ PDF created successfully!")
        return str(pdf_path)
        
    except ImportError:
        print("⚠ pywin32 not installed - PDF conversion skipped")
        print(f"✓ Word document created: {docx_path}")
        print("\nTo enable automatic PDF conversion:")
        print("  pip install pywin32")
        return str(docx_path)
        
    except Exception as e:
        print(f"⚠ Could not convert to PDF: {e}")
        print(f"✓ Word document created: {docx_path}")
        return str(docx_path)


def create_resume_from_file(excel_path, output_dir, generate_pdf=True, generate_docx=True):
    """
    Create resume from Excel file for API use
    
    Args:
        excel_path: Path to Excel template file
        output_dir: Directory to save output files
        generate_pdf: Whether to generate PDF (default True)
        generate_docx: Whether to generate DOCX (default True)
    
    Returns:
        dict with success status and file paths
    """
    result = {
        'success': False,
        'name': '',
        'docx_path': None,
        'pdf_path': None,
        'error': None
    }
    
    try:
        print(f"Loading Excel file: {excel_path}")
        wb = load_workbook(excel_path, data_only=True)
        
        # Read defaults
        defaults = read_defaults(wb)
        
        print("Creating Word document...")
        doc = Document()
        
        # Set up styles
        style = doc.styles['Normal']
        try:
            style.font.name = 'Times New Roman'
        except:
            pass
        style.font.size = Pt(10)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 1.0
        
        # Page setup
        section = doc.sections[0]
        section.page_height = Inches(11)
        section.page_width = Inches(8.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        
        # Read section order
        section_order = read_section_order(wb)
        
        # Pre-read all data
        contact = read_contact_info(wb['Contact_Info'])
        
        summary = None
        summary_name = 'Professional Summary'
        if 'Summary' in wb.sheetnames:
            summary = read_summary(wb['Summary'])
            custom = read_custom_section_name(wb['Summary'])
            if custom:
                summary_name = custom
        
        highlights = []
        highlights_name = 'Highlights'
        if 'Highlights' in wb.sheetnames:
            highlights = read_bullet_section(wb['Highlights'])
            custom = read_custom_section_name(wb['Highlights'])
            if custom:
                highlights_name = custom
        
        achievements = []
        achievements_name = 'Achievements'
        if 'Achievements' in wb.sheetnames:
            achievements = read_bullet_section(wb['Achievements'])
            custom = read_custom_section_name(wb['Achievements'])
            if custom:
                achievements_name = custom
        
        licenses = []
        licenses_name = 'Licenses & Certifications'
        if 'Licenses_Certifications' in wb.sheetnames:
            licenses = read_bullet_section(wb['Licenses_Certifications'])
            custom = read_custom_section_name(wb['Licenses_Certifications'])
            if custom:
                licenses_name = custom
        
        memberships = []
        memberships_name = 'Memberships'
        if 'Memberships' in wb.sheetnames:
            memberships = read_bullet_section(wb['Memberships'])
            custom = read_custom_section_name(wb['Memberships'])
            if custom:
                memberships_name = custom
        
        awards = []
        awards_name = 'Awards'
        if 'Awards' in wb.sheetnames:
            awards = read_bullet_section(wb['Awards'])
            custom = read_custom_section_name(wb['Awards'])
            if custom:
                awards_name = custom
        
        projects = []
        projects_name = 'Projects'
        if 'Projects' in wb.sheetnames:
            projects = read_bullet_section(wb['Projects'])
            custom = read_custom_section_name(wb['Projects'])
            if custom:
                projects_name = custom
        
        custom_section = []
        custom_section_name = 'Additional Information'
        if 'Custom_Section' in wb.sheetnames:
            custom_section = read_bullet_section(wb['Custom_Section'])
            custom = read_custom_section_name(wb['Custom_Section'])
            if custom:
                custom_section_name = custom
        
        skills_name = 'Skills'
        if 'Skills' in wb.sheetnames:
            custom = read_custom_section_name(wb['Skills'])
            if custom:
                skills_name = custom
        
        work_roles_name = 'Professional Experience'
        if 'Work_Roles' in wb.sheetnames:
            custom = read_custom_section_name(wb['Work_Roles'])
            if custom:
                work_roles_name = custom
        
        work_roles = read_work_roles(wb['Work_Roles'])
        
        education_name = 'Education'
        if 'Education' in wb.sheetnames:
            custom = read_custom_section_name(wb['Education'])
            if custom:
                education_name = custom
        
        education = read_education(wb['Education'])
        
        # Build resume
        for section_name in section_order:
            if section_name == 'Contact_Info':
                build_contact_header(doc, contact, defaults)
            elif section_name == 'Summary':
                build_summary(doc, summary, defaults, summary_name)
            elif section_name == 'Highlights':
                build_bullet_section(doc, highlights_name, highlights, defaults)
            elif section_name == 'Work_Roles':
                build_work_experience(doc, work_roles, wb['Work_Experience'], defaults, work_roles_name)
            elif section_name == 'Education':
                build_education(doc, education, defaults, education_name)
            elif section_name == 'Licenses_Certifications':
                build_bullet_section(doc, licenses_name, licenses, defaults)
            elif section_name == 'Skills':
                build_skills(doc, wb['Skills'], defaults, skills_name)
            elif section_name == 'Achievements':
                build_bullet_section(doc, achievements_name, achievements, defaults)
            elif section_name == 'Memberships':
                build_bullet_section(doc, memberships_name, memberships, defaults)
            elif section_name == 'Awards':
                build_bullet_section(doc, awards_name, awards, defaults)
            elif section_name == 'Projects':
                build_bullet_section(doc, projects_name, projects, defaults)
            elif section_name == 'Custom_Section':
                build_bullet_section(doc, custom_section_name, custom_section, defaults)
        
        # Create filename
        full_name = contact.get('Name', 'Resume')
        name_parts = full_name.split()
        
        if len(name_parts) >= 2:
            first_name = name_parts[0]
            last_name = name_parts[-1]
            base_filename = f"{first_name}_{last_name}"
        else:
            base_filename = full_name.replace(' ', '_') if full_name else "Resume"
        
        timestamp = datetime.now().strftime("%Y-%m-%d_%H%M%S")
        filename = f"{base_filename}_{timestamp}"
        
        output_path = Path(output_dir)
        docx_path = output_path / f"{filename}.docx"
        pdf_path = output_path / f"{filename}.pdf"
        
        result['name'] = base_filename
        
        # Save DOCX (always if we need PDF, even temporarily)
        if generate_docx or generate_pdf:
            doc.save(str(docx_path))
            if generate_docx:
                result['docx_path'] = str(docx_path)
        
        # Convert to PDF using CloudConvert API
        if generate_pdf:
            try:
                import requests
                import time
                import os
                
                API_KEY = os.getenv('CLOUDCONVERT_API_KEY')
                
                if not API_KEY:
                    print("⚠️ CloudConvert API key not configured")
                    result['pdf_path'] = None
                else:
                    print("🔄 Starting PDF conversion via CloudConvert...")
                    
                    job_response = requests.post(
                        'https://api.cloudconvert.com/v2/jobs',
                        headers={
                            'Authorization': f'Bearer {API_KEY}',
                            'Content-Type': 'application/json'
                        },
                        json={
                            'tasks': {
                                'import-docx': {
                                    'operation': 'import/upload'
                                },
                                'convert-to-pdf': {
                                    'operation': 'convert',
                                    'input': 'import-docx',
                                    'output_format': 'pdf'
                                },
                                'export-pdf': {
                                    'operation': 'export/url',
                                    'input': 'convert-to-pdf'
                                }
                            }
                        }
                    )
                    
                    if job_response.status_code != 201:
                        print(f"❌ Failed to create job: {job_response.text}")
                        result['pdf_path'] = None
                    else:
                        job = job_response.json()['data']
                        upload_task = job['tasks'][0]
                        
                        print("📤 Uploading DOCX...")
                        with open(docx_path, 'rb') as f:
                            upload_response = requests.post(
                                upload_task['result']['form']['url'],
                                data=upload_task['result']['form']['parameters'],
                                files={'file': f}
                            )
                        
                        if upload_response.status_code not in [200, 201]:
                            print(f"❌ Upload failed: {upload_response.text}")
                            result['pdf_path'] = None
                        else:
                            print("⏳ Converting to PDF...")
                            job_id = job['id']
                            
                            for attempt in range(30):
                                status_response = requests.get(
                                    f'https://api.cloudconvert.com/v2/jobs/{job_id}',
                                    headers={'Authorization': f'Bearer {API_KEY}'}
                                )
                                
                                status_data = status_response.json()['data']
                                
                                if status_data['status'] == 'finished':
                                    print("📥 Downloading PDF...")
                                    export_task = [t for t in status_data['tasks'] if t['name'] == 'export-pdf'][0]
                                    pdf_url = export_task['result']['files'][0]['url']
                                    
                                    pdf_response = requests.get(pdf_url)
                                    
                                    with open(pdf_path, 'wb') as f:
                                        f.write(pdf_response.content)
                                    
                                    result['pdf_path'] = str(pdf_path)
                                    print(f"✅ PDF created: {pdf_path}")
                                    break
                                
                                elif status_data['status'] == 'error':
                                    print(f"❌ Conversion error: {status_data}")
                                    result['pdf_path'] = None
                                    break
                                
                                time.sleep(1)
                            else:
                                print("❌ Conversion timed out")
                                result['pdf_path'] = None
                        
                        # Clean up temp DOCX if user only wanted PDF
                        if result.get('pdf_path') and not generate_docx:
                            try:
                                os.remove(docx_path)
                                print(f"🗑️ Removed temporary DOCX")
                            except:
                                pass
                        
            except Exception as e:
                print(f"❌ PDF conversion exception: {e}")
                import traceback
                traceback.print_exc()
                result['pdf_path'] = None
        
        result['success'] = True
        return result
        
    except Exception as e:
        result['error'] = str(e)
        print(f"Error creating resume: {e}")
        import traceback
        traceback.print_exc()
        return result


def preview_resume_content(excel_path):
    """
    Preview resume content without generating files
    """
    try:
        wb = load_workbook(excel_path, data_only=True)
        
        # Get basic info
        contact = read_contact_info(wb['Contact_Info'])
        section_order = read_section_order(wb)
        
        # Count content
        work_roles = read_work_roles(wb['Work_Roles'])
        
        return {
            'success': True,
            'name': contact.get('Name', 'Unknown'),
            'sections_included': section_order,
            'total_work_roles': len(work_roles)
        }
    except Exception as e:
        return {
            'success': False,
            'error': str(e)
        }


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python resume_generator.py <excel_file>")
        sys.exit(1)
    
    excel_file = sys.argv[1]
    result = create_resume(excel_file)
    print(f"\n✓ Resume created: {result}")
